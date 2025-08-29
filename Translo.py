import os
import json
import uuid
import docx
import PyPDF2
import feedparser
import random
import math
import io
from typing import Tuple, Optional, Dict, Any
from flask import Flask, request, jsonify, render_template_string, send_from_directory, Response
import requests
from bs4 import BeautifulSoup, NavigableString
from openai import OpenAI
from flask_cors import CORS
# 新增: 用于生成 PDF 和 DOCX
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.utils import simpleSplit

# --- 1. Flask 应用初始化 ---
app = Flask(__name__)
CORS(app)
# 创建一个临时文件夹用于存放翻译后的文件
TEMP_FOLDER = 'translated_files'
os.makedirs(TEMP_FOLDER, exist_ok=True)
app.config['TEMP_FOLDER'] = TEMP_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB 文件大小限制

# --- 字体设置 (已修复) ---
# 为了解决PDF中文乱码，需要一个支持中文的字体文件。
# 推荐使用 "Noto Sans SC"。请下载 NotoSansSC-Regular.otf 并将其放置在与此脚本相同的目录中。
# 下载地址: https://fonts.google.com/noto/specimen/Noto+Sans+SC

# 【修复】: 构造字体文件的绝对路径，确保任何情况下都能找到它
# `os.path.abspath(__file__)` 获取此脚本的绝对路径
# `os.path.dirname(...)` 获取脚本所在的目录
# `os.path.join(...)` 将目录和文件名安全地拼接起来
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FONT_PATH = os.path.join(BASE_DIR, 'NotoSansSC-Regular.otf')

FONT_NAME = "NotoSansSC"
try:
    # 【修复】: 使用构造好的绝对路径来注册字体
    pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))
    print(f"成功加载字体: {FONT_PATH}")
except Exception as e:
    # 打印更详细的错误信息，方便调试
    print(f"警告: 在路径 '{FONT_PATH}' 未找到字体文件。PDF中的中文可能无法正确显示。将回退到默认字体。错误详情: {e}")
    FONT_NAME = "Helvetica" # 回退

# --- 2. 模型配置 (新增价格和限制) ---
# 价格单位: 美元 / 每百万 token
MODEL_CONFIG = {
    "openai": {
        "base_url": None, "model": "gpt-4o-mini", "name": "OpenAI GPT-4o-mini",
        "limit_chars": 100000, "price_input_pm": 0.15, "price_output_pm": 0.60
    },
    "deepseek": {
        "base_url": "https://api.deepseek.com/v1", "model": "deepseek-chat", "name": "DeepSeek Chat",
        "limit_chars": 100000, "price_input_pm": 0.14, "price_output_pm": 0.28 # 价格为示例
    },
    "grok": {
        "base_url": "https://api.groq.com/openai/v1", "model": "llama3-8b-8192", "name": "Grok (Llama3 on Groq)",
        "limit_chars": 30000, "price_input_pm": 0.05, "price_output_pm": 0.05 # 价格为示例
    },
    "tongyi": {
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1", "model": "qwen-turbo", "name": "Tongyi Qwen (通义千问)",
        "limit_chars": 30000, "price_input_pm": 0.1, "price_output_pm": 0.1 # 价格为示例, 假设价格
    }
}

# --- 3. 迭代翻译核心 ---
def iterative_translation_client(original_text: str, target_language: str, api_key: str, model_choice: str, refinement_prompt: str, quality: str, glossary: str) -> str:
    config = MODEL_CONFIG.get(model_choice)
    if not config:
        raise ValueError(f"无效的模型选择: {model_choice}")

    print(f"使用模型进行翻译: {config['name']}, 质量模式: {quality}")
    client = OpenAI(api_key=api_key, base_url=config["base_url"])

    # --- 步骤 1: 初始翻译 ---
    initial_prompt = f"请将以下文本翻译成“{target_language}”。请只返回翻译后的文本。\n\n原文：\n\"{original_text}\""
    try:
        initial_response = client.chat.completions.create(
            model=config["model"],
            messages=[
                {"role": "system", "content": "你是一个专业的翻译助手。"},
                {"role": "user", "content": initial_prompt}
            ]
        )
        draft_translation = initial_response.choices[0].message.content.strip()
    except Exception as e:
        print(f"调用 {config['name']} API 时出错 (初始翻译): {e}")
        return f"[{config['name']} 初始翻译失败]: {e}"

    if quality == 'fast':
        return draft_translation

    # --- 步骤 2: 反思与优化 (仅高质量模式) ---
    print("进行第二步：反思与优化...")
    
    final_refinement_prompt = refinement_prompt or "请仔细审校这份翻译初稿，检查其准确性、流畅度和语调。修正所有错误，并改进措辞，使其听起来对母语者来说既自然又专业。"
    if glossary:
        final_refinement_prompt += f"\n\n请务必遵守以下术语表：\n{glossary}"

    refinement_system_prompt = "你是一位顶级的翻译审校专家。你的任务是根据用户提供的优化指令，审校并改进一份翻译初稿，使其更完美。请只输出最终的、优化后的译文。"
    refinement_user_prompt = f"""
    请根据以下信息，优化翻译初稿：
    - **原始文本**: "{original_text}"
    - **翻译初稿**: "{draft_translation}"
    - **优化指令 (请严格遵守)**: "{final_refinement_prompt}"
    请输出最终优化后的译文:
    """
    try:
        refined_response = client.chat.completions.create(
            model=config["model"],
            messages=[
                {"role": "system", "content": refinement_system_prompt},
                {"role": "user", "content": refinement_user_prompt}
            ]
        )
        final_translation = refined_response.choices[0].message.content.strip()
        return final_translation
    except Exception as e:
        print(f"调用 {config['name']} API 时出错 (优化翻译): {e}")
        return f"[{config['name']} 优化翻译失败]: {e}"

def universal_translation_client(payload: Dict[str, Any], api_key: str, model_choice: str, is_json: bool) -> str:
    # ... (此函数无需修改)
    config = MODEL_CONFIG.get(model_choice)
    if not config:
        raise ValueError(f"无效的模型选择: {model_choice}")
    
    try:
        client = OpenAI(api_key=api_key, base_url=config["base_url"])
        request_params = {"model": config["model"], "messages": payload["messages"]}
        if is_json:
            request_params["response_format"] = {"type": "json_object"}
        response = client.chat.completions.create(**request_params)
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"调用 {config['name']} API 时出错: {e}")
        if is_json:
            error_map = {key: f"[{config['name']} 翻译失败]" for key in payload["original_keys"]}
            return json.dumps(error_map)
        else:
            return f"[{config['name']} 翻译失败]: {e}"

def iterative_batch_translation(texts_to_translate: dict, target_language: str, api_key: str, model_choice: str, refinement_prompt: str, quality: str, glossary: str) -> dict:
    # ... (此函数无需修改)
    initial_prompt = f"""
    你是一个专业的翻译引擎。请将下面 JSON 对象中的字符串值翻译成“{target_language}”。
    保持原始的 JSON 结构，你的输出必须是一个格式正确的 JSON 对象。
    原始 JSON:
    {json.dumps(texts_to_translate, ensure_ascii=False, indent=2)}
    """
    payload = {
        "messages": [{"role": "system", "content": "你是一个专门用于输出 JSON 格式翻译结果的助手。"}, {"role": "user", "content": initial_prompt}],
        "original_keys": list(texts_to_translate.keys())
    }
    try:
        draft_json_str = universal_translation_client(payload, api_key, model_choice, is_json=True)
        draft_translations = json.loads(draft_json_str)
    except json.JSONDecodeError:
        print("API返回的不是有效的JSON，对批量翻译中止。")
        return {key: "[翻译失败：模型返回格式错误]" for key in texts_to_translate.keys()}


    if quality == 'fast':
        return draft_translations

    print("进行第二步：批量反思与优化...")
    
    final_refinement_prompt = refinement_prompt or "请仔细审校这份翻译初稿，检查其准确性、流畅度和语调。修正所有错误，并改进措辞，使其听起来对母语者来说既自然又专业。"
    if glossary:
        final_refinement_prompt += f"\n\n请务必遵守以下术语表：\n{glossary}"

    refinement_system_prompt = "你是一位顶级的翻译审校专家。你的任务是根据用户提供的优化指令，审校并改进一个包含翻译初稿的 JSON 对象。请保持 JSON 结构，只优化其中的译文，然后输出最终的、优化后的 JSON 对象。"
    refinement_user_prompt = f"""
    请根据以下信息，优化 JSON 对象中的翻译初稿：
    - **原始文本 JSON**: {json.dumps(texts_to_translate, ensure_ascii=False, indent=2)}
    - **翻译初稿 JSON**: {json.dumps(draft_translations, ensure_ascii=False, indent=2)}
    - **优化指令 (请严格遵守所有译文)**: "{final_refinement_prompt}"
    请输出最终优化后的 JSON 对象:
    """
    
    refinement_payload = {
        "messages": [{"role": "system", "content": refinement_system_prompt}, {"role": "user", "content": refinement_user_prompt}],
        "original_keys": list(texts_to_translate.keys())
    }
    try:
        final_json_str = universal_translation_client(refinement_payload, api_key, model_choice, is_json=True)
        return json.loads(final_json_str)
    except json.JSONDecodeError:
        print("API在优化步骤返回的不是有效的JSON。")
        return {key: "[翻译失败：模型返回格式错误]" for key in texts_to_translate.keys()}


# --- 4. Agentic Workflow (网页翻译) ---
def fetch_and_parse(url: str) -> Tuple[Optional[str], Optional[dict]]:
    # ... (此函数无需修改)
    print(f"Agent 1: 开始抓取和解析 URL -> {url}")
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        response.encoding = response.apparent_encoding
    except requests.RequestException as e:
        print(f"错误：抓取 URL {url} 失败: {e}")
        return None, None

    soup = BeautifulSoup(response.text, 'html.parser')
    for element in soup(["script", "style"]):
        element.decompose()

    text_map = {}
    for element in soup.find_all(string=True):
        if isinstance(element, NavigableString) and element.strip() and element.parent.name not in ['script', 'style', 'head', 'title', 'meta', '[document]']:
            text = element.strip()
            if text:
                placeholder_id = f"trans_{uuid.uuid4().hex[:8]}"
                text_map[placeholder_id] = text
                element.replace_with(f"{{{placeholder_id}}}")
    
    print(f"Agent 1: 解析完成，提取到 {len(text_map)} 条待翻译文本。")
    return str(soup), text_map

def reconstruct_page(html_template: str, translated_map: dict) -> str:
    # ... (此函数无需修改)
    print("Agent 3: 开始重构页面...")
    final_html = html_template
    for placeholder_id, translated_text in translated_map.items():
        final_html = final_html.replace(f"{{{placeholder_id}}}", translated_text)
    print("Agent 3: 页面重构完成。")
    return final_html

# --- 5. 文本与文件翻译功能 ---
def parse_file(file_stream, filename: str) -> Optional[str]:
    """从内存中的文件流提取文本"""
    if filename.endswith('.txt'):
        return file_stream.read().decode('utf-8', errors='ignore')
    elif filename.endswith('.docx'):
        doc = docx.Document(file_stream)
        return "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(file_stream)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            print(f"PDF 解析失败: {e}")
            return None
    return None

def create_translated_pdf(text: str, file_path: str):
    """使用 ReportLab 创建一个包含翻译文本的 PDF"""
    # *** 修复点2: 优化PDF生成逻辑 ***
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    style = styles['Normal']
    style.fontName = FONT_NAME
    style.fontSize = 10
    style.leading = 12

    # 将换行符替换为 <br/> 标签，以便在单个 Paragraph 中正确处理换行
    # 这比为每一行创建一个新的 Paragraph 效果更好
    text_with_breaks = text.replace('\n', '<br/>')
    story = [Paragraph(text_with_breaks, style)]
    
    doc.build(story)

def create_translated_docx(text: str, file_path: str):
    """创建一个包含翻译文本的 DOCX 文件"""
    doc = docx.Document()
    for para in text.split('\n'):
        doc.add_paragraph(para)
    doc.save(file_path)

def get_lang_code(language_name: str) -> str:
    """将语言全名映射到短代码"""
    lang_map = {
        "Simplified Chinese": "zh", "English": "en", "Japanese": "ja",
        "Korean": "ko", "Spanish": "es", "French": "fr",
        "German": "de", "Russian": "ru"
    }
    return lang_map.get(language_name, "trans")

# --- 6. API 端点 ---
NEWS_FEEDS = [
    "https://news.google.com/rss?hl=zh-CN&gl=CN&ceid=CN:zh-Hans", # Google News (China)
    "http://feeds.reuters.com/reuters/worldNews", # Reuters World News
    "http://feeds.bbci.co.uk/news/world/rss.xml" # BBC World News
]

@app.route('/get-trends')
def get_trends():
    """从多个源随机获取实时热点新闻标题"""
    feed_url = random.choice(NEWS_FEEDS)
    print(f"从以下地址获取新闻: {feed_url}")
    try:
        feed = feedparser.parse(feed_url)
        headlines = [entry.title for entry in feed.entries[:5]]
        return jsonify(headlines)
    except Exception as e:
        print(f"获取 RSS feed 失败: {e}")
        return jsonify([]), 500

@app.route('/translate-url', methods=['POST'])
def handle_url_translation():
    data = request.json
    url = data.get('url')
    target_language = data.get('language', '中文')
    api_key = data.get('api_key')
    model_choice = data.get('model', 'openai')
    refinement_prompt = data.get('refinement_prompt', '')
    glossary = data.get('glossary', '')
    quality = data.get('quality', 'high')

    if not url or not api_key:
        return jsonify({"error": "URL 和 API Key 都是必需的"}), 400

    html_template, texts_to_translate = fetch_and_parse(url)
    if html_template is None:
        return jsonify({"error": "无法抓取或解析该 URL。"}), 500
    if not texts_to_translate:
        return jsonify({"error": "在该页面上未找到可翻译的文本。"}), 400

    translated_texts = iterative_batch_translation(texts_to_translate, target_language, api_key, model_choice, refinement_prompt, quality, glossary)
    final_html = reconstruct_page(html_template, translated_texts)
    return final_html

@app.route('/translate-text', methods=['POST'])
def handle_text_translation():
    data = request.json
    text = data.get('text')
    target_language = data.get('language', '中文')
    api_key = data.get('api_key')
    model_choice = data.get('model', 'openai')
    refinement_prompt = data.get('refinement_prompt', '')
    glossary = data.get('glossary', '')
    quality = data.get('quality', 'high')

    if not text or not api_key:
        return jsonify({"error": "待翻译文本和 API Key 都是必需的"}), 400

    translated_text = iterative_translation_client(text, target_language, api_key, model_choice, refinement_prompt, quality, glossary)
    return jsonify({"translated_text": translated_text})

@app.route('/translate-file', methods=['POST'])
def handle_file_translation():
    if 'file_id' not in request.form:
         return jsonify({"error": "缺少文件ID"}), 400

    file_id = request.form.get('file_id')
    original_filename = request.form.get('original_filename')
    temp_file_path = os.path.join(app.config['TEMP_FOLDER'], file_id)

    if not os.path.exists(temp_file_path):
        return jsonify({"error": "文件不存在或已过期"}), 404

    api_key = request.form.get('api_key')
    target_language = request.form.get('language', '中文')
    model_choice = request.form.get('model', 'openai')
    refinement_prompt = request.form.get('refinement_prompt', '')
    glossary = request.form.get('glossary', '')
    quality = request.form.get('quality', 'high')

    if not api_key:
        return jsonify({"error": "API Key 是必需的"}), 400

    def generate_translation():
        try:
            with open(temp_file_path, 'rb') as f:
                file_stream = io.BytesIO(f.read())
            
            text_content = parse_file(file_stream, original_filename)
            
            # 删除临时上传文件
            os.remove(temp_file_path)

            if text_content is None:
                raise ValueError("不支持的文件类型或文件解析失败。")

            chunk_size = 3000 
            text_chunks = [text_content[i:i + chunk_size] for i in range(0, len(text_content), chunk_size)]
            total_chunks = len(text_chunks)
            
            translated_chunks = []
            for i, chunk in enumerate(text_chunks):
                progress_data = json.dumps({"progress": i + 1, "total": total_chunks})
                yield f"data: {progress_data}\n\n"
                if chunk.strip():
                    translated_chunk = iterative_translation_client(chunk, target_language, api_key, model_choice, refinement_prompt, quality, glossary)
                    translated_chunks.append(translated_chunk)
            
            final_translated_text = "".join(translated_chunks)
            preview_text = (final_translated_text[:1500] + '...') if len(final_translated_text) > 1500 else final_translated_text
            
            original_basename, original_ext = os.path.splitext(original_filename)
            lang_code = get_lang_code(target_language)
            output_ext = '.pdf' if original_ext == '.pdf' else ('.docx' if original_ext == '.docx' else '.txt')
            download_filename = f"{original_basename}_{lang_code}_translated{output_ext}"
            file_path = os.path.join(app.config['TEMP_FOLDER'], download_filename)

            if output_ext == '.pdf':
                create_translated_pdf(final_translated_text, file_path)
            elif output_ext == '.docx':
                create_translated_docx(final_translated_text, file_path)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(final_translated_text)
            
            result_data = json.dumps({"preview_text": preview_text, "download_filename": download_filename})
            yield f"data: {result_data}\n\n"

        except Exception as e:
            # 记录服务器端的完整错误信息
            print(f"文件翻译过程中发生错误: {e}")
            # 向客户端发送一个通用、安全的错误信息
            error_data = json.dumps({"error": f"服务器处理文件时出错，请查看后台日志了解详情。"})
            yield f"data: {error_data}\n\n"

    return Response(generate_translation(), mimetype='text/event-stream')


@app.route('/upload-for-estimation', methods=['POST'])
def upload_for_estimation():
    if 'file' not in request.files:
        return jsonify({"error": "未找到文件部分"}), 400
    files = request.files.getlist('file')
    model_choice = request.form.get('model', 'openai')
    quality = request.form.get('quality', 'high')

    config = MODEL_CONFIG.get(model_choice)
    if not config:
        return jsonify({"error": "无效的模型选择"}), 400

    total_char_count = 0
    file_ids = []
    original_filenames = []

    for file in files:
        file_id = str(uuid.uuid4())
        temp_path = os.path.join(app.config['TEMP_FOLDER'], file_id)
        file.save(temp_path)

        with open(temp_path, 'rb') as f_read:
            file_stream = io.BytesIO(f_read.read())
        
        text_content = parse_file(file_stream, file.filename)
        if text_content:
            total_char_count += len(text_content)
            file_ids.append(file_id)
            original_filenames.append(file.filename)
        else:
            os.remove(temp_path) # 如果文件无效，则删除

    if total_char_count == 0:
        return jsonify({"error": "无法从文件中读取内容进行估算。"}), 400

    estimated_tokens = total_char_count / 2.5
    cost_per_million = config['price_input_pm'] + config['price_output_pm']
    estimated_cost = (estimated_tokens / 1_000_000) * cost_per_million
    if quality == 'high':
        estimated_cost *= 2 

    chunk_size = 3000
    num_chunks = math.ceil(total_char_count / chunk_size)
    time_per_chunk = 10 if quality == 'high' else 5
    estimated_time_seconds = num_chunks * time_per_chunk

    return jsonify({
        "file_id": file_ids[0], # 简化为处理单个文件
        "original_filename": original_filenames[0],
        "char_count": total_char_count,
        "estimated_cost_usd": round(estimated_cost, 4),
        "estimated_time_seconds": estimated_time_seconds
    })


@app.route('/download/<filename>')
def download_file(filename):
    """提供翻译后文件的下载"""
    try:
        return send_from_directory(app.config['TEMP_FOLDER'], filename, as_attachment=True)
    except FileNotFoundError:
        return jsonify({"error": "文件未找到或已过期。"}), 404

# --- 7. 前端界面 ---
@app.route('/')
def index():
    model_options_html = ""
    for key, config in MODEL_CONFIG.items():
        model_options_html += f'<option value="{key}" data-limit="{config["limit_chars"]}">{config["name"]}</option>'

    html_content = f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Agentic Workflow 多功能翻译器</title>
        <style>
            :root {{
                --primary-color: #007bff; --primary-hover: #0056b3; --background-color: #f8f9fa;
                --card-background: #ffffff; --text-color: #343a40; --border-color: #dee2e6;
                --error-color: #dc3545; --light-gray: #f1f3f5; --fast-color: #28a745; --fast-hover: #218838;
                --warning-color: #ffc107;
            }}
            body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; background-color: var(--background-color); color: var(--text-color); display: flex; flex-direction: column; align-items: center; margin: 0; padding: 1rem; min-height: 100vh; }}
            .container {{ width: 100%; max-width: 800px; background: var(--card-background); padding: 2rem; border-radius: 12px; box-shadow: 0 6px 20px rgba(0,0,0,0.08); box-sizing: border-box; }}
            .main-container {{ max-width: 1200px; }}
            h1 {{ color: var(--primary-color); text-align: center; margin-bottom: 0.5rem; }}
            p.subtitle {{ text-align: center; margin-top: 0; color: #6c757d; margin-bottom: 2rem; }}
            .mode-selector {{ display: flex; justify-content: center; margin-bottom: 1.5rem; background: var(--light-gray); border-radius: 8px; padding: 5px; }}
            .mode-selector button {{ flex: 1; padding: 0.7rem; border: none; background: transparent; cursor: pointer; font-size: 1rem; font-weight: 600; border-radius: 6px; transition: background-color 0.3s, color 0.3s; }}
            .mode-selector button.active {{ background-color: var(--primary-color); color: white; }}
            .form-group {{ margin-bottom: 1.5rem; }}
            label {{ display: block; margin-bottom: 0.5rem; font-weight: 600; }}
            input, select, textarea {{ width: 100%; padding: 0.8rem 1rem; border: 1px solid var(--border-color); border-radius: 8px; box-sizing: border-box; font-size: 1rem; font-family: inherit; transition: border-color 0.2s, box-shadow 0.2s; }}
            textarea {{ resize: vertical; }}
            input:focus, select:focus, textarea:focus {{ outline: none; border-color: var(--primary-color); box-shadow: 0 0 0 3px rgba(0, 123, 255, 0.25); }}
            
            /* --- 按钮布局优化 --- */
            #action-buttons-container {{ 
                margin-top: 2rem;
                display: flex;
                justify-content: center;
                width: 100%;
            }}
            .translation-buttons {{ 
                display: flex;
                gap: 1.5rem;
                width: 100%;
                max-width: 500px;
            }}
            .main-translate-btn {{ 
                flex: 1; /* 让按钮平分空间 */
                padding: 1rem;
                color: white;
                border: none;
                border-radius: 8px;
                font-size: 1.1rem;
                font-weight: 600;
                cursor: pointer;
                transition: background-color 0.3s, transform 0.1s;
            }}
            .fast-btn {{ background-color: var(--fast-color); }}
            .fast-btn:hover:not(:disabled) {{ background-color: var(--fast-hover); }}
            .high-btn {{ background-color: var(--primary-color); }}
            .high-btn:hover:not(:disabled) {{ background-color: var(--primary-hover); }}
            .stop-btn {{ background-color: var(--error-color); }}
            .stop-btn:hover:not(:disabled) {{ background-color: #c82333; }}
            
            /* 优化后的加载器样式 */
            #loader {{
                display: none; /* 默认隐藏 */
                position: fixed;
                top: 0;
                left: 0;
                width: 100%;
                height: 100%;
                background-color: rgba(255, 255, 255, 0.4); /* 更透明的背景 */
                backdrop-filter: blur(4px);
                z-index: 1000;
                justify-content: center;
                align-items: center;
                flex-direction: column;
                text-align: center;
            }}
            #hide-loader-btn {{
                position: absolute;
                top: 20px;
                right: 20px;
                background: rgba(0, 0, 0, 0.1);
                border: none;
                font-size: 1.5rem;
                color: #333;
                cursor: pointer;
                line-height: 1;
                width: 40px;
                height: 40px;
                border-radius: 50%;
                display: flex;
                justify-content: center;
                align-items: center;
                transition: background-color 0.2s;
            }}
            #hide-loader-btn:hover {{
                background: rgba(0, 0, 0, 0.2);
            }}
            .loader-content svg {{
                width: 80px;
                height: 80px;
                margin-bottom: 1rem;
            }}
            .loader-content .pen {{
                fill: none;
                stroke: var(--primary-color);
                stroke-width: 5;
                stroke-linecap: round;
                stroke-linejoin: round;
                animation: draw 3s linear infinite;
            }}
            @keyframes draw {{
                to {{
                    stroke-dashoffset: 0;
                }}
            }}
            #loader-tip {{
                font-size: 1.1rem;
                color: var(--text-color);
                max-width: 80%;
                margin-top: 1rem;
                transition: opacity 0.5s ease-in-out;
            }}
            #stop-translation-btn-loader {{
                margin-top: 2rem;
                padding: 0.8rem 1.5rem;
                font-size: 1rem;
                width: auto;
                flex: none; /* *** 修复点1: 防止按钮在加载动画中被拉伸 *** */
            }}

            .result-area {{ margin-top: 2rem; width: 100%; }}
            .text-comparison-area {{ display: grid; grid-template-columns: 1fr 1fr; gap: 1rem; height: 300px; }}
            .text-box {{ display: flex; flex-direction: column; background: var(--light-gray); border: 1px solid var(--border-color); border-radius: 8px; }}
            .text-box h3 {{ margin-top: 0; padding: 0.5rem 1rem; border-bottom: 1px solid var(--border-color); }}
            .text-box textarea, .text-box div {{ flex-grow: 1; padding: 1rem; border: none; background: transparent; white-space: pre-wrap; overflow-y: auto; }}
            .text-box textarea:focus {{ outline: none; box-shadow: none; }}
            #result-frame {{ display: none; width: 100%; height: 70vh; border: 1px solid var(--border-color); border-radius: 8px; }}
            .error-message {{ color: var(--error-color); margin-top: 1rem; text-align: center; font-weight: 500; }}
            .hidden {{ display: none !important; }}
            #file-upload-area label.file-label {{ background-color: var(--light-gray); color: var(--text-color); padding: 2rem; border-radius: 8px; text-align: center; cursor: pointer; display: block; border: 2px dashed var(--border-color); transition: background-color 0.2s, border-color 0.2s; }}
            #file-upload-area label.dragover {{ background-color: #e9ecef; border-color: var(--primary-color); }}
            #file-upload-area input[type="file"] {{ display: none; }}
            #file-name {{ margin-top: 1rem; text-align: center; color: #6c757d; }}
            #file-result-area {{ display: none; }}
            .file-warning {{ color: var(--warning-color); font-weight: bold; margin-top: 1rem; }}
            .api-note {{ font-size: 0.85rem; color: #6c757d; margin-top: 0.5rem; }}

            /* 历史记录样式 */
            #history-interface {{ max-height: 60vh; overflow-y: auto; }}
            .history-item {{ border: 1px solid var(--border-color); border-radius: 8px; padding: 1rem; margin-bottom: 1rem; }}
            .history-item p {{ margin: 0.5rem 0; white-space: pre-wrap; word-wrap: break-word; }}
            .history-item .meta {{ font-size: 0.85rem; color: #6c757d; border-bottom: 1px solid var(--border-color); padding-bottom: 0.5rem; margin-bottom: 0.5rem; }}
            .history-item .content {{ max-height: 100px; overflow-y: auto; background: #fdfdfd; padding: 0.5rem; border-radius: 4px; }}
            .history-actions {{ margin-top: 1rem; display: flex; gap: 0.5rem; }}
            .history-actions button {{ font-size: 0.8rem; padding: 0.3rem 0.6rem; cursor: pointer; }}
            .history-settings {{ display: flex; justify-content: space-between; align-items: center; margin-top: 1rem; padding-top: 1rem; border-top: 1px solid var(--border-color); }}
            .history-settings .form-group {{ display: flex; align-items: center; gap: 1rem; margin-bottom: 0; }}
            #history-limit {{ width: 80px; }}
            .privacy-notice {{ background-color: var(--light-gray); padding: 1rem; border-radius: 8px; text-align: center; color: #6c757d; margin-bottom: 1rem; }}

            /* 新增布局样式 */
            #controls-area-grid {{
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 2rem;
            }}
            @media (max-width: 768px) {{
                #controls-area-grid {{
                    grid-template-columns: 1fr;
                }}
            }}

        </style>
    </head>
    <body>
        <div class="container main-container">
            <h1>多功能翻译器</h1>
            <p class="subtitle">支持文本、网页和文件翻译</p>
            
            <div class="mode-selector">
                <button id="mode-btn-text" class="active">文本翻译</button>
                <button id="mode-btn-url">网页翻译</button>
                <button id="mode-btn-file">文件翻译</button>
                <button id="mode-btn-history">历史记录</button>
            </div>

            <!-- 文本翻译界面 -->
            <div id="text-translation-interface">
                <div class="text-comparison-area">
                    <div class="text-box"><h3>原文</h3><textarea id="text-input" placeholder="在此输入需要翻译的内容..."></textarea></div>
                    <div class="text-box"><h3>译文</h3><textarea id="text-output" readonly placeholder="翻译结果将显示在此处..."></textarea></div>
                </div>
            </div>

            <!-- 网页翻译界面 -->
            <div id="url-translation-interface" class="hidden">
                <div class="form-group"><label for="url-input">网页 URL</label><input type="url" id="url-input" placeholder="https://example.com"></div>
            </div>

            <!-- 文件翻译界面 -->
            <div id="file-translation-interface" class="hidden">
                <div class="form-group" id="file-upload-area">
                    <label for="file-input" class="file-label">点击或拖拽文件到此处 (.txt, .pdf, .docx)</label>
                    <input type="file" id="file-input" accept=".txt,.pdf,.docx" multiple>
                    <p id="file-name"></p>
                    <p id="file-warning-message" class="file-warning hidden"></p>
                </div>
            </div>

            <!-- 历史记录界面 -->
            <div id="history-interface" class="hidden">
                <p class="privacy-notice">您的所有翻译记录均安全地存储在您自己的浏览器中，绝不会上传至服务器。</p>
                <div id="history-list"></div>
                <div class="history-settings">
                    <div class="form-group">
                        <label for="history-limit">保存记录条数:</label>
                        <input type="number" id="history-limit" min="1" max="100" value="10">
                    </div>
                    <div>
                        <button id="export-history-btn" class="main-translate-btn fast-btn" style="width: auto; padding: 0.5rem 1rem;">导出历史</button>
                        <button id="clear-history-btn" class="main-translate-btn stop-btn" style="width: auto; padding: 0.5rem 1rem;">清除历史</button>
                    </div>
                </div>
            </div>
            
            <div id="controls-area">
                <div id="controls-area-grid">
                    <div id="left-controls">
                        <div class="form-group">
                            <label for="glossary">术语表 (可选, 每行一条, 格式: 原文:译文)</label>
                            <textarea id="glossary" placeholder="例如:&#10;Agentic Workflow: 智能体工作流&#10;Large Language Model: 大语言模型" style="height: 120px;"></textarea>
                        </div>
                        <div class="form-group">
                            <label for="refinement-prompt">优化指令 (可选)</label>
                            <textarea id="refinement-prompt" placeholder="例如：语气正式、使用英式英语..." style="height: 120px;"></textarea>
                        </div>
                    </div>
                    <div id="right-controls">
                        <div class="form-group"><label for="model-select">选择 AI 模型</label><select id="model-select">{model_options_html}</select></div>
                        <div class="form-group"><label for="language">目标语言</label><select id="language"><option value="Simplified Chinese">简体中文</option><option value="English">英语</option><option value="Japanese">日语</option><option value="Korean">韩语</option><option value="Spanish">西班牙语</option><option value="French">法语</option><option value="German">德语</option><option value="Russian">俄语</option></select></div>
                        <div class="form-group">
                            <label for="api_key">你的大模型 API Key</label>
                            <input type="password" id="api_key" placeholder="请输入所选模型的有效 API Key" required>
                            <p class="api-note">API密钥和翻译内容仅用于本次请求，服务器不会保存或记录任何敏感信息。</p>
                        </div>
                    </div>
                </div>
                <!-- *** 优化点2: 按钮容器移动到此处并居中 *** -->
                <div id="action-buttons-container">
                    <div class="translation-buttons" id="start-buttons">
                        <button id="translate-fast-btn" class="main-translate-btn fast-btn">快速翻译</button>
                        <button id="translate-high-btn" class="main-translate-btn high-btn">高质量翻译</button>
                    </div>
                    <button id="large-file-btn" class="main-translate-btn high-btn hidden">翻译大型文件 (估算成本)</button>
                    <button id="stop-translation-btn-main" class="main-translate-btn stop-btn hidden">停止翻译</button>
                </div>
            </div>
            <div id="error-msg" class="error-message"></div>

            <!-- 文件翻译结果预览 -->
            <div id="file-result-area" class="result-area hidden">
                <h3>翻译预览</h3>
                <div class="text-box" style="height: 200px;"><div id="file-preview-content"></div></div>
            </div>
        </div>

        <!-- 网页翻译结果 -->
        <div class="result-area container main-container">
            <iframe id="result-frame" title="翻译结果"></iframe>
        </div>

        <!-- 全屏加载器 -->
        <div id="loader">
            <button id="hide-loader-btn">&times;</button>
            <div class="loader-content">
                <svg viewBox="0 0 100 100">
                    <path class="pen" d="M20,80 L40,20 L60,80 L80,20" stroke-dasharray="200" stroke-dashoffset="200"></path>
                </svg>
                <h3 id="loader-title">正在进行高质量翻译...</h3>
                <p id="loader-tip">你知道吗？全世界大约有7000种仍在使用中的语言。</p>
                <div id="progress-container" class="hidden" style="width: 80%; max-width: 400px; margin-top: 1rem;">
                    <p id="progress-text"></p>
                    <progress id="progress-bar" value="0" max="100" style="width: 100%;"></progress>
                </div>
            </div>
            <button id="stop-translation-btn-loader" class="main-translate-btn stop-btn">停止翻译</button>
        </div>

        <script>
            const modeButtons = {{ text: document.getElementById('mode-btn-text'), url: document.getElementById('mode-btn-url'), file: document.getElementById('mode-btn-file'), history: document.getElementById('mode-btn-history') }};
            const interfaces = {{ text: document.getElementById('text-translation-interface'), url: document.getElementById('url-translation-interface'), file: document.getElementById('file-translation-interface'), history: document.getElementById('history-interface') }};
            const controlsArea = document.getElementById('controls-area');
            
            const startButtonsContainer = document.getElementById('start-buttons');
            const largeFileBtn = document.getElementById('large-file-btn');
            const stopButtonMain = document.getElementById('stop-translation-btn-main');
            const stopButtonLoader = document.getElementById('stop-translation-btn-loader');

            const loader = document.getElementById('loader');
            const hideLoaderBtn = document.getElementById('hide-loader-btn');
            const errorMsg = document.getElementById('error-msg');
            
            const textInput = document.getElementById('text-input');
            const textOutput = document.getElementById('text-output');
            const urlInput = document.getElementById('url-input');
            const fileInput = document.getElementById('file-input');
            const fileUploadArea = document.getElementById('file-upload-area');
            const fileLabel = fileUploadArea.querySelector('label');
            const fileNameDisplay = document.getElementById('file-name');
            const fileWarningMessage = document.getElementById('file-warning-message');

            const langSelect = document.getElementById('language');
            const apiKeyInput = document.getElementById('api_key');
            const modelSelect = document.getElementById('model-select');
            const refinementPromptInput = document.getElementById('refinement-prompt');
            const glossaryInput = document.getElementById('glossary');
            
            const fileResultArea = document.getElementById('file-result-area');
            const filePreviewContent = document.getElementById('file-preview-content');
            const urlResultFrame = document.getElementById('result-frame');

            const historyList = document.getElementById('history-list');
            const clearHistoryBtn = document.getElementById('clear-history-btn');
            const exportHistoryBtn = document.getElementById('export-history-btn');
            const historyLimitInput = document.getElementById('history-limit');

            const progressContainer = document.getElementById('progress-container');
            const progressBar = document.getElementById('progress-bar');
            const progressText = document.getElementById('progress-text');
            
            const loaderTip = document.getElementById('loader-tip');
            let tipInterval = null;
            let trendsCache = [];
            let lastTrendFetch = 0;
            let currentAbortController = null;
            let debounceTimer = null;
            let progressInterval = null;

            const staticTips = [
                "你知道吗？全世界大约有7000种仍在使用中的语言。",
                "“机器人”（Robot）一词源于捷克语，意为“强制劳动”。",
                "莎Shakespeare在他的作品中创造了超过1700个新单词。",
                "“翻译”不仅仅是文字替换，更是跨文化的沟通艺术。",
                "芬兰语和匈牙利语是欧洲少数不属于印欧语系的语言。",
                "巴布亚新几内亚是世界上语言最多样化的国家，拥有超过800种语言。",
                "最长的英文单词有45个字母，是一种肺病的名称。",
                "“Deja vu”（既视感）是一个法语词，意思是“已经见过”。",
                "在手语中，不同的国家和地区也有自己的“方言”。",
                "字母“J”是最后一个被添加到英文字母表中的字母。"
            ];

            // --- History Functions ---
            const HISTORY_KEY = 'translationHistory';
            const HISTORY_LIMIT_KEY = 'translationHistoryLimit';

            function getHistoryLimit() {{
                return parseInt(localStorage.getItem(HISTORY_LIMIT_KEY) || '10', 10);
            }}

            function setHistoryLimit(limit) {{
                localStorage.setItem(HISTORY_LIMIT_KEY, limit);
                // Trim history if new limit is smaller
                const history = getHistory();
                if (history.length > limit) {{
                    const trimmedHistory = history.slice(0, limit);
                    localStorage.setItem(HISTORY_KEY, JSON.stringify(trimmedHistory));
                    renderHistory();
                }}
            }}

            function getHistory() {{
                return JSON.parse(localStorage.getItem(HISTORY_KEY) || '[]');
            }}
            function saveHistory(entry) {{
                const limit = getHistoryLimit();
                const history = getHistory();
                history.unshift(entry); // Add to the beginning
                if (history.length > limit) {{
                    history.pop();
                }}
                localStorage.setItem(HISTORY_KEY, JSON.stringify(history));
            }}
            function clearHistory() {{
                if (confirm('确定要清除所有历史记录吗？此操作不可恢复。')) {{
                    localStorage.removeItem(HISTORY_KEY);
                    renderHistory();
                }}
            }}
            function exportHistory() {{
                const history = getHistory();
                const jsonString = JSON.stringify(history, null, 2);
                const blob = new Blob([jsonString], {{ type: 'application/json' }});
                const url = URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = 'translation_history.json';
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                URL.revokeObjectURL(url);
            }}
            function renderHistory() {{
                const history = getHistory();
                if (history.length === 0) {{
                    historyList.innerHTML = '<p>暂无历史记录。</p>';
                    return;
                }}
                historyList.innerHTML = history.map(item => `
                    <div class="history-item">
                        <p class="meta">
                            <strong>模式:</strong> ${{item.mode}} | <strong>模型:</strong> ${{item.model}} | <strong>质量:</strong> ${{item.quality}} | <strong>语言:</strong> ${{item.language}} | <em>${{new Date(item.timestamp).toLocaleString()}}</em>
                        </p>
                        <p><strong>原文:</strong></p>
                        <div class="content">${{item.original}}</div>
                        <p><strong>译文:</strong></p>
                        <div class="content">${{item.translated}}</div>
                        <div class="history-actions">
                            <button onclick="copyToClipboard('${{btoa(unescape(encodeURIComponent(item.original)))}}', true)">复制原文</button>
                            <button onclick="copyToClipboard('${{btoa(unescape(encodeURIComponent(item.translated)))}}', true)">复制译文</button>
                        </div>
                    </div>
                `).join('');
            }}
            function copyToClipboard(base64Text, isBase64 = false) {{
                const text = isBase64 ? decodeURIComponent(escape(atob(base64Text))) : base64Text;
                const textarea = document.createElement('textarea');
                textarea.value = text;
                document.body.appendChild(textarea);
                textarea.select();
                document.execCommand('copy');
                document.body.removeChild(textarea);
                alert('已复制到剪贴板！');
            }}

            async function fetchTrends() {{
                const now = Date.now();
                if (trendsCache.length > 0 && now - lastTrendFetch < 300000) {{
                    return; // 使用缓存
                }}
                try {{
                    const controller = new AbortController();
                    const timeoutId = setTimeout(() => controller.abort(), 5000); // 5秒超时
                    const response = await fetch('/get-trends', {{ signal: controller.signal }});
                    clearTimeout(timeoutId);

                    if (response.ok) {{
                        const headlines = await response.json();
                        if (headlines && headlines.length > 0) {{
                            trendsCache = headlines.map(h => `实时热点：${{h}}`);
                            lastTrendFetch = now;
                        }}
                    }}
                }} catch (error) {{
                    console.error("获取热点新闻失败 (可能超时):", error);
                    trendsCache = []; // 清空缓存以备下次尝试
                }}
            }}

            function showNextTip() {{
                loaderTip.style.opacity = 0;
                
                const combinedTips = [...staticTips, ...trendsCache];
                const randomIndex = Math.floor(Math.random() * combinedTips.length);
                
                setTimeout(() => {{
                    loaderTip.textContent = combinedTips[randomIndex];
                    loaderTip.style.opacity = 1;
                }}, 500);
            }}

            let currentMode = 'text';

            function switchMode(mode) {{
                currentMode = mode;
                errorMsg.textContent = '';
                urlResultFrame.style.display = 'none';
                fileResultArea.classList.add('hidden');
                
                for (const key in modeButtons) {{
                    modeButtons[key].classList.toggle('active', key === mode);
                }}
                for (const key in interfaces) {{
                     interfaces[key].classList.toggle('hidden', key !== mode);
                }}

                controlsArea.classList.toggle('hidden', mode === 'history');
                if (mode === 'history') {{
                    renderHistory();
                }}
            }}

            Object.keys(modeButtons).forEach(key => {{
                modeButtons[key].addEventListener('click', () => switchMode(key));
            }});
            
            fileInput.addEventListener('change', () => {{
                handleFileSelection(fileInput.files);
            }});

            ['dragenter', 'dragover', 'dragleave', 'drop'].forEach(eventName => {{
                fileUploadArea.addEventListener(eventName, (e) => {{ e.preventDefault(); e.stopPropagation(); }}, false);
            }});
            ['dragenter', 'dragover'].forEach(eventName => {{
                fileUploadArea.addEventListener(eventName, () => fileLabel.classList.add('dragover'), false);
            }});
            ['dragleave', 'drop'].forEach(eventName => {{
                fileUploadArea.addEventListener(eventName, () => fileLabel.classList.remove('dragover'), false);
            }});
            fileUploadArea.addEventListener('drop', (e) => {{
                fileInput.files = e.dataTransfer.files;
                handleFileSelection(fileInput.files);
            }}, false);

            function handleFileSelection(files) {{
                fileWarningMessage.classList.add('hidden');
                startButtonsContainer.classList.remove('hidden');
                largeFileBtn.classList.add('hidden');
                
                if (files.length === 0) {{
                    fileNameDisplay.textContent = '';
                    return;
                }}
                if (files.length > 5) {{
                    fileWarningMessage.textContent = '错误：一次最多只能上传5个文件。';
                    fileWarningMessage.classList.remove('hidden');
                    fileInput.value = ''; // 清空选择
                    return;
                }}

                let totalSize = 0;
                let fileNames = [];
                for(const file of files) {{
                    totalSize += file.size;
                    fileNames.push(file.name);
                }}

                if (totalSize > 20 * 1024 * 1024) {{
                    fileWarningMessage.textContent = '错误：文件总大小不能超过20MB。';
                    fileWarningMessage.classList.remove('hidden');
                    fileInput.value = '';
                    return;
                }}
                
                fileNameDisplay.textContent = fileNames.join(', ');

                // 检查字符数
                const reader = new FileReader();
                reader.onload = function(e) {{
                    const text = e.target.result;
                    const charCount = text.length;
                    const selectedModel = modelSelect.options[modelSelect.selectedIndex];
                    const limit = parseInt(selectedModel.dataset.limit, 10);
                    if(charCount > limit) {{
                        fileWarningMessage.textContent = `注意：文件字符数 (${{charCount}}) 超过单次翻译限制 (${{limit}})。请使用下方的“翻译大型文件”按钮。`;
                        fileWarningMessage.classList.remove('hidden');
                        startButtonsContainer.classList.add('hidden');
                        largeFileBtn.classList.remove('hidden');
                    }}
                }};
                // 只读取第一个文件来估算
                reader.readAsText(files[0]);
            }}
            
            function stopTranslation() {{
                if (currentAbortController) {{
                    currentAbortController.abort();
                }}
            }}

            hideLoaderBtn.addEventListener('click', () => {{
                loader.style.display = 'none';
            }});
            stopButtonMain.addEventListener('click', stopTranslation);
            stopButtonLoader.addEventListener('click', stopTranslation);
            clearHistoryBtn.addEventListener('click', clearHistory);
            exportHistoryBtn.addEventListener('click', exportHistory);
            historyLimitInput.addEventListener('change', () => {{
                setHistoryLimit(historyLimitInput.value);
            }});


            function resetUI(message = '') {{
                loader.style.display = 'none';
                if (tipInterval) clearInterval(tipInterval);
                if (progressInterval) clearInterval(progressInterval);
                progressContainer.classList.add('hidden');
                startButtonsContainer.classList.remove('hidden');
                stopButtonMain.classList.add('hidden');
                errorMsg.textContent = message;
                currentAbortController = null;
            }}

            async function triggerTranslation(quality, isLargeFile = false) {{
                if (currentAbortController) return; // 防止重复点击

                const language = langSelect.value;
                const apiKey = apiKeyInput.value.trim();
                const model = modelSelect.value;
                const refinementPrompt = refinementPromptInput.value.trim();
                const glossary = glossaryInput.value.trim();

                if (!apiKey) {{ errorMsg.textContent = '请输入您的 API Key。'; return; }}

                startButtonsContainer.classList.add('hidden');
                largeFileBtn.classList.add('hidden');
                stopButtonMain.classList.remove('hidden');
                
                if (quality === 'high') {{
                    loader.style.display = 'flex';
                    // 立即显示一条静态提示，然后在后台获取新闻
                    showNextTip(); 
                    fetchTrends(); 
                    tipInterval = setInterval(showNextTip, 7000); // 延长提示时间
                }}

                errorMsg.textContent = '';
                urlResultFrame.style.display = 'none';
                fileResultArea.classList.add('hidden');
                currentAbortController = new AbortController();

                try {{
                    if (isLargeFile) {{
                        await handleFileTranslation(language, apiKey, model, refinementPrompt, quality, glossary, isLargeFile);
                    }} else {{
                        if (currentMode === 'text') await handleTextTranslation(language, apiKey, model, refinementPrompt, quality, glossary);
                        else if (currentMode === 'url') await handleUrlTranslation(language, apiKey, model, refinementPrompt, quality, glossary);
                        else if (currentMode === 'file') await handleFileTranslation(language, apiKey, model, refinementPrompt, quality, glossary, isLargeFile);
                    }}
                    if (progressInterval) progressBar.value = 100;
                    resetUI();
                }} catch (error) {{
                    if (error.name === 'AbortError') {{
                        console.log('Fetch aborted by user.');
                        resetUI('翻译已取消。');
                        if (currentMode === 'text') {{
                            textOutput.value = '翻译已取消。';
                        }}
                    }} else {{
                        console.error('翻译失败:', error);
                        resetUI(`翻译失败: ${{error.message}}`);
                    }}
                }}
            }}

            document.getElementById('translate-fast-btn').addEventListener('click', () => triggerTranslation('fast'));
            document.getElementById('translate-high-btn').addEventListener('click', () => triggerTranslation('high'));
            largeFileBtn.addEventListener('click', () => handleLargeFileTranslation());

            async function handleLargeFileTranslation() {{
                const files = fileInput.files;
                if (files.length === 0) return;

                const formData = new FormData();
                for(const file of files) {{
                    formData.append('file', file);
                }}
                formData.append('model', modelSelect.value);
                formData.append('quality', 'high'); // 大型文件默认高质量

                const estimateResponse = await fetch('/upload-for-estimation', {{ method: 'POST', body: formData }});
                if (!estimateResponse.ok) {{
                    const err = await estimateResponse.json();
                    errorMsg.textContent = `估算失败: ${{err.error}}`;
                    return;
                }}
                const estimate = await estimateResponse.json();
                
                const confirmation = confirm(
                    `翻译此大型文件 (${{estimate.char_count}} 字符) 详情如下:\\n` +
                    `- 预计费用: $${{estimate.estimated_cost_usd}} USD\\n` +
                    `- 预计耗时: 约 ${{Math.ceil(estimate.estimated_time_seconds / 60)}} 分钟\\n\\n` +
                    `是否继续？`
                );

                if (confirmation) {{
                    triggerTranslation('high', estimate);
                }}
            }}


            async function handleTextTranslation(language, apiKey, model, refinementPrompt, quality, glossary) {{
                const text = textInput.value.trim();
                if (!text) {{ throw new Error('请输入需要翻译的文本。'); }}
                
                textOutput.value = "正在翻译...";
                const response = await fetch('/translate-text', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ text, language, api_key: apiKey, model, refinement_prompt: refinementPrompt, quality, glossary }}),
                    signal: currentAbortController.signal
                }});
                if (!response.ok) {{ const err = await response.json(); throw new Error(err.error || '服务器返回错误'); }}
                
                const result = await response.json();
                textOutput.value = result.translated_text;
                
                saveHistory({{
                    mode: '文本',
                    original: text,
                    translated: result.translated_text,
                    model: modelSelect.options[modelSelect.selectedIndex].text,
                    quality: quality === 'high' ? '高质量' : '快速',
                    language: langSelect.options[langSelect.selectedIndex].text,
                    timestamp: new Date().toISOString()
                }});
            }}

            async function handleUrlTranslation(language, apiKey, model, refinementPrompt, quality, glossary) {{
                const url = urlInput.value.trim();
                if (!url) {{ throw new Error('请输入有效的网页 URL。'); }}

                const response = await fetch('/translate-url', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ url, language, api_key: apiKey, model, refinement_prompt: refinementPrompt, quality, glossary }}),
                    signal: currentAbortController.signal
                }});
                if (!response.ok) {{ const err = await response.json(); throw new Error(err.error || '服务器返回错误'); }}
                
                const translatedHtml = await response.text();
                urlResultFrame.srcdoc = translatedHtml;
                urlResultFrame.style.display = 'block';

                saveHistory({{
                    mode: '网页',
                    original: url,
                    translated: '网页内容已翻译，请在下方查看。',
                    model: modelSelect.options[modelSelect.selectedIndex].text,
                    quality: quality === 'high' ? '高质量' : '快速',
                    language: langSelect.options[langSelect.selectedIndex].text,
                    timestamp: new Date().toISOString()
                }});
            }}

            // *** 优化点1: 修复文件翻译流式处理的 Bug ***
            async function handleFileTranslation(language, apiKey, model, refinementPrompt, quality, glossary, largeFileEstimate = false) {{
                const formData = new FormData();
                
                if(largeFileEstimate) {{
                    formData.append('file_id', largeFileEstimate.file_id);
                    formData.append('original_filename', largeFileEstimate.original_filename);
                }} else {{
                    if (fileInput.files.length === 0) {{ throw new Error('请选择一个文件。'); }}
                    // 在此版本中，我们简化为一次处理一个文件以匹配后端逻辑
                    formData.append('file_id', 'placeholder'); // 后端需要这个字段
                    formData.append('original_filename', fileInput.files[0].name);
                }}

                formData.append('language', language);
                formData.append('api_key', apiKey);
                formData.append('model', model);
                formData.append('refinement_prompt', refinementPrompt);
                formData.append('quality', quality);
                formData.append('glossary', glossary);

                progressContainer.classList.remove('hidden');

                const response = await fetch('/translate-file', {{ method: 'POST', body: formData, signal: currentAbortController.signal }});
                
                if (!response.ok) {{
                    const err = await response.json();
                    throw new Error(err.error || '服务器返回错误');
                }}

                const reader = response.body.getReader();
                const decoder = new TextDecoder();
                let buffer = ''; // 使用 buffer 来存储不完整的消息

                while(true) {{
                    const {{ value, done }} = await reader.read();
                    if (done) break;

                    buffer += decoder.decode(value, {{ stream: true }}); // 将新数据块附加到缓冲区
                    let boundary = buffer.indexOf('\\n\\n'); // 查找消息分隔符

                    while (boundary !== -1) {{
                        const message = buffer.substring(0, boundary); // 提取完整的消息
                        buffer = buffer.substring(boundary + 2); // 从缓冲区中移除已处理的消息

                        if (message.startsWith('data: ')) {{
                            try {{
                                const dataStr = message.substring(6);
                                const data = JSON.parse(dataStr);

                                if (data.progress) {{
                                    progressBar.value = data.progress;
                                    progressBar.max = data.total;
                                    progressText.textContent = `正在翻译 ${{data.progress}} / ${{data.total}} 块...`;
                                }} else if (data.download_filename) {{
                                    filePreviewContent.textContent = data.preview_text;
                                    fileResultArea.classList.remove('hidden');

                                    const link = document.createElement('a');
                                    link.href = `/download/${{data.download_filename}}`;
                                    link.setAttribute('download', data.download_filename);
                                    document.body.appendChild(link);
                                    link.click();
                                    document.body.removeChild(link);

                                    saveHistory({{
                                        mode: '文件',
                                        original: `文件名: ${{largeFileEstimate ? largeFileEstimate.original_filename : fileInput.files[0].name}}`,
                                        translated: '文件已翻译并下载。',
                                        model: modelSelect.options[modelSelect.selectedIndex].text,
                                        quality: quality === 'high' ? '高质量' : '快速',
                                        language: langSelect.options[langSelect.selectedIndex].text,
                                        timestamp: new Date().toISOString()
                                    }});
                                }} else if (data.error) {{
                                    throw new Error(data.error);
                                }}
                            }} catch (e) {{
                                console.error("解析服务器消息失败:", e);
                                console.error("原始消息:", message);
                                throw new Error("无法解析来自服务器的数据。");
                            }}
                        }}
                        boundary = buffer.indexOf('\\n\\n'); // 在剩余的缓冲区中查找下一个消息
                    }}
                }}
                progressContainer.classList.add('hidden');
            }}
            
            textInput.addEventListener('keyup', () => {{
                clearTimeout(debounceTimer);
                debounceTimer = setTimeout(() => {{
                    if (textInput.value.trim().length > 0) {{
                        triggerTranslation('fast');
                    }}
                }}, 1500);
            }});

            // Initialize
            historyLimitInput.value = getHistoryLimit();
            switchMode('text');
        </script>
    </body>
    </html>
    """
    return render_template_string(html_content)

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5001, debug=True)